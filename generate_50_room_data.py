import os
import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches  # 用於設置字體大小和頁面邊距
from docx.enum.text import WD_LINE_SPACING  # 用於設置行距
from docx.enum.style import WD_STYLE_TYPE  # 用於檢查樣式類型
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
from datetime import datetime
from lunar_python import Solar, Lunar  # 用於農曆轉換
from tkcalendar import Calendar  # 添加日曆選擇器

# ============================
# 輔助函式
# ----------------------------

def set_paragraph_font(paragraph, font_name="標楷體", font_size=27):
    """設置段落的字體和字體大小，避免覆蓋字符間距"""
    paragraph.style.font.name = font_name
    paragraph.style.font.size = Pt(font_size)
    for run in paragraph.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 確保中文使用指定字體
        run.font.size = Pt(font_size)
        # 不設置字符間距，保留模板的設置

def set_paragraph_spacing(paragraph):
    """設置段落的行距和間距，確保一致（包括空白行）"""
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 0.4  # 行距調整為 0.4 倍
    paragraph.paragraph_format.space_before = Pt(0)  # 段前間距為 0
    paragraph.paragraph_format.space_after = Pt(0)   # 段後間距為 0
    paragraph.paragraph_format.keep_together = True  # 確保不分頁
    paragraph.paragraph_format.keep_with_next = True  # 與下段同頁

def parse_ancestors(tablet1: str) -> str:
    """解析牌位1列中的祖先姓名，將名字分配到兩行，每行最多兩個名字，用空格分隔"""
    if pd.isna(tablet1):
        return ""
    # 按換行符分割多個祖先名稱
    ancestors = [line.strip() for line in tablet1.split("\n") if line.strip()]
    if not ancestors:
        return ""
    
    # 將名字分配到兩行，每行最多兩個名字
    if len(ancestors) <= 2:
        # 如果名字數量少於或等於 2，直接用空格分隔
        return " ".join(ancestors)
    else:
        # 將名字分為兩組：前兩個名字一行，後面的名字另一行
        first_line = " ".join(ancestors[:2])  # 前兩個名字
        second_line = " ".join(ancestors[2:])  # 後面的名字
        return f"{first_line}\n{second_line}" if second_line else first_line

def format_address(address: str) -> str:
    """格式化地址，確保換行後每一行前面有一個 TAB 空格"""
    max_length = 20  # 每行最多 20 個字符
    if len(address) > max_length:
        parts = [address[i:i + max_length] for i in range(0, len(address), max_length)]
        result = parts[0]  # 第一行不加 TAB
        for part in parts[1:]:  # 第二行及之後加 TAB
            result += f"\n\t{part}"
        return result
    return address

def convert_to_chinese_date(date_str: str) -> str:
    """將YYYY-MM-DD格式的日期轉換為中國農曆格式"""
    try:
        date_obj = datetime.strptime(date_str, "%Y-%m-%d")
        solar = Solar.fromYmdHms(date_obj.year, date_obj.month, date_obj.day, 0, 0, 0)
        lunar = solar.getLunar()
        year_name = lunar.getYearInGanZhi()
        month_names = ["正", "二", "三", "四", "五", "六", "七", "八", "九", "十", "十一", "十二"]
        day_names = ["初一", "初二", "初三", "初四", "初五", "初六", "初七", "初八", "初九", "初十", 
                     "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十",
                     "廿一", "廿二", "廿三", "廿四", "廿五", "廿六", "廿七", "廿八", "廿九", "三十", "卅一"]
        return f"{year_name}年 {month_names[lunar.getMonth() - 1]}月 {day_names[lunar.getDay() - 1]}日"
    except ValueError as e:
        print(f"日期轉換錯誤: {e}")
        return "無效日期"

def generate_word_file(template_path: str, row: pd.Series, date_str: str, output_dir: str, output_text: tk.Text) -> None:
    """為單行數據生成Word文件（基於模板）"""
    try:
        # 載入模板文件
        doc = Document(template_path)

        # 設置頁面邊距（0.3 英寸）
        section = doc.sections[0]
        section.left_margin = Inches(0.3)
        section.right_margin = Inches(0.3)
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.3)

        # 獲取數據
        ancestors = parse_ancestors(row.get("牌位1", ""))
        name = row.get("姓名", "未知姓名").strip()
        address = row.get("您的住址", "未知地址")
        chinese_date = convert_to_chinese_date(date_str)

        # 格式化地址
        address = format_address(address)

        # 檢查內容長度，必要時截斷
        if len(address) > 50:
            address = address[:47] + "..."
        ancestors_list = parse_ancestors(row.get("牌位1", "")).split("\n")
        if len(ancestors_list) > 2:  # 限制「牌位1」最多兩行
            ancestors = "\n".join(ancestors_list[:2])

        # 動態計算字體大小（降低閾值，加快縮減速度）
        total_length = len(ancestors) + len(address) + len(name) + len(chinese_date)
        font_size = 27
        if total_length > 30:  # 降低閾值，從 40 調整為 30
            font_size = max(10, 27 - (total_length - 30) // 5)  # 每增加 5 個字符減小 1 點字體大小，最小 10

        print(f"計算出的字體大小: {font_size} (總長度: {total_length})")  # 調試輸出

        # 遍歷文檔中的所有段落，查找並替換標記
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if "{牌位1}" in text:
                paragraph.text = text.replace("{牌位1}", ancestors)
            if "{姓名}" in text:
                paragraph.text = text.replace("{姓名}", name)
            if "{地址}" in text:
                paragraph.text = text.replace("{地址}", address)
            if "{日期}" in text:
                paragraph.text = text.replace("{日期}", chinese_date)

        # 遍歷文檔中的所有段落，設置字體和格式（包括空白行）
        for paragraph in doc.paragraphs:
            set_paragraph_font(paragraph, font_size=font_size)
            set_paragraph_spacing(paragraph)  # 統一設置行距和間距

        # 額外檢查：確保樣式級別的行距也被設置（只處理段落樣式）
        for style in doc.styles:
            if style.type == WD_STYLE_TYPE.PARAGRAPH:  # 只處理段落樣式
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                style.paragraph_format.line_spacing = 0.4  # 行距調整為 0.4 倍
                style.paragraph_format.space_before = Pt(0)
                style.paragraph_format.space_after = Pt(0)

        # 保存文件
        output_file = os.path.join(output_dir, f"{name}.docx")
        doc.save(output_file)
        print(f"已生成: {output_file}")
        output_text.insert(tk.END, f"已生成: {output_file}\n")
        output_text.see(tk.END)

    except Exception as e:
        error_msg = f"生成Word文件時發生錯誤 ({row.get('姓名', '未知姓名')}): {str(e)}"
        print(error_msg)
        output_text.insert(tk.END, f"{error_msg}\n")
        output_text.see(tk.END)

def process_excel_file(file_path: str, template_path: str, date_str: str, output_text: tk.Text) -> None:
    """處理Excel文件並生成Word文件"""
    try:
        # 讀取 Excel 文件，指定 header=4（從第 5 行開始作為欄位名稱）
        df = pd.read_excel(file_path, header=4)

        # 打印讀取到的欄位名稱
        print("讀取到的欄位名稱:", list(df.columns))

        # 打印前幾行數據
        print("前幾行數據:\n", df.head())

        # 檢查必要欄位是否存在
        required_columns = ["姓名", "牌位1", "您的住址"]
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Excel文件缺少必要的欄位: {', '.join(missing_columns)}")

        # 清理數據：過濾掉無效行
        df = df[df["姓名"].notna() & df["姓名"].apply(lambda x: isinstance(x, str) and x.strip() != "")]

        # 創建輸出資料夾
        output_dir = os.path.join(os.path.dirname(file_path), "Output") if os.path.dirname(file_path) else "Output"
        os.makedirs(output_dir, exist_ok=True)

        # 檢查 output_dir 是否可寫入
        if not os.access(output_dir, os.W_OK):
            raise PermissionError(f"無法寫入資料夾: {output_dir}")

        # 逐行生成 Word 文件
        for index, row in df.iterrows():
            generate_word_file(template_path, row, date_str, output_dir, output_text)

    except Exception as e:
        error_msg = f"處理Excel文件時發生錯誤: {str(e)}"
        print(error_msg)
        messagebox.showerror("錯誤", error_msg)

# ============================
# 主應用程式類別
# ----------------------------

class App:
    """Excel到Word轉換器的主應用程式類別"""

    def __init__(self, root: tk.Tk) -> None:
        self.root = root
        self._setup_window()
        self._create_widgets()
        self.file_path = ""
        self.template_path = ""
        self.selected_date = "2025-03-24"

    def _setup_window(self) -> None:
        """設置主視窗的基本屬性 (暗色介面)"""
        self.root.title("Excel到Word轉換器")
        self.root.geometry("600x600")
        self.root.configure(bg='#2E2E2E')

        style = ttk.Style()
        style.theme_use('clam')
        dark_bg = '#2E2E2E'
        dark_fg = '#FFFFFF'
        entry_bg = '#3E3E3E'
        for widget in ['TButton', 'TProgressbar', 'TFrame', 'TLabel']:
            style.configure(widget, background=dark_bg, foreground=dark_fg, relief=tk.FLAT, font=("微軟正黑體", 10))
        style.configure("TEntry", background=entry_bg, foreground=dark_fg, fieldbackground=entry_bg, relief=tk.FLAT, font=("微軟正黑體", 10))
        style.configure("TCombobox", background=dark_bg, foreground="black", relief=tk.FLAT, font=("微軟正黑體", 10))
        style.map("TButton", background=[('active', '#3E3E3E')])

    def _create_widgets(self) -> None:
        """創建和配置所有UI元件"""
        # 選擇 Excel 文件
        input_frame = ttk.Frame(self.root, style='TFrame')
        input_frame.place(relx=0.5, rely=0.08, anchor=tk.CENTER)
        label = tk.Label(input_frame, text="選擇Excel文件:", font=("微軟正黑體", 12), bg='#2E2E2E', fg='#FFFFFF')
        label.pack(side=tk.LEFT, padx=5)
        self.select_button = tk.Button(input_frame, text="選擇文件", command=self.select_file,
                                       bg="#4CAF50", fg="white", font=("微軟正黑體", 10), relief=tk.FLAT)
        self.select_button.pack(side=tk.LEFT, padx=5)

        self.path_label = tk.Label(self.root, text="", font=("微軟正黑體", 10), bg='#2E2E2E', fg='#FFFFFF')
        self.path_label.place(relx=0.5, rely=0.16, anchor=tk.CENTER)

        # 選擇模板文件
        template_frame = ttk.Frame(self.root, style='TFrame')
        template_frame.place(relx=0.5, rely=0.24, anchor=tk.CENTER)
        template_label = tk.Label(template_frame, text="選擇模板文件:", font=("微軟正黑體", 12), bg='#2E2E2E', fg='#FFFFFF')
        template_label.pack(side=tk.LEFT, padx=5)
        self.select_template_button = tk.Button(template_frame, text="選擇模板", command=self.select_template,
                                                bg="#4CAF50", fg="white", font=("微軟正黑體", 10), relief=tk.FLAT)
        self.select_template_button.pack(side=tk.LEFT, padx=5)

        self.template_label = tk.Label(self.root, text="", font=("微軟正黑體", 10), bg='#2E2E2E', fg='#FFFFFF')
        self.template_label.place(relx=0.5, rely=0.32, anchor=tk.CENTER)

        # 選擇日期
        date_frame = ttk.Frame(self.root, style='TFrame')
        date_frame.place(relx=0.5, rely=0.40, anchor=tk.CENTER)
        date_label = tk.Label(date_frame, text="選擇日期:", font=("微軟正黑體", 12), bg='#2E2E2E', fg='#FFFFFF')
        date_label.pack(side=tk.LEFT, padx=5)

        self.date_entry = ttk.Entry(date_frame, width=15, style="TEntry")
        self.date_entry.pack(side=tk.LEFT, padx=5)
        self.date_entry.insert(0, "2025-03-24")
        self.date_entry.bind("<FocusOut>", self.validate_date_entry)

        self.date_button = tk.Button(date_frame, text="選擇", command=self.open_calendar,
                                     bg="#2196F3", fg="white", font=("微軟正黑體", 10), relief=tk.FLAT)
        self.date_button.pack(side=tk.LEFT, padx=5)

        # 開始生成按鈕
        button_frame = ttk.Frame(self.root, style='TFrame')
        button_frame.place(relx=0.5, rely=0.50, anchor=tk.CENTER)
        self.start_button = tk.Button(button_frame, text="開始生成", command=self.start_conversion,
                                      state=tk.DISABLED, bg="#2196F3", fg="white", font=("微軟正黑體", 10), relief=tk.FLAT)
        self.start_button.pack(pady=5)

        # 輸出結果區域
        output_frame = ttk.Frame(self.root, style='TFrame')
        output_frame.place(relx=0.5, rely=0.75, anchor=tk.CENTER, width=500, height=200)

        output_label = tk.Label(output_frame, text="文件輸出結果:", font=("微軟正黑體", 12), bg='#2E2E2E', fg='#FFFFFF')
        output_label.pack(anchor=tk.W, padx=5)

        self.output_text = tk.Text(output_frame, height=10, width=60, bg='#3E3E3E', fg='#FFFFFF', font=("微軟正黑體", 10), relief=tk.FLAT)
        scrollbar = ttk.Scrollbar(output_frame, orient=tk.VERTICAL, command=self.output_text.yview)
        self.output_text.configure(yscrollcommand=scrollbar.set)
        self.output_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

    def select_file(self) -> None:
        """選擇Excel文件"""
        self.file_path = filedialog.askopenfilename(filetypes=[("Excel檔案", "*.xlsx *.xls")])
        if self.file_path and self.template_path:
            self.start_button.config(state=tk.NORMAL)
        self.path_label.config(text=f"當前文件: {self.file_path}")

    def select_template(self) -> None:
        """選擇模板文件"""
        self.template_path = filedialog.askopenfilename(filetypes=[("Word檔案", "*.docx *.doc")])
        if self.file_path and self.template_path:
            self.start_button.config(state=tk.NORMAL)
        self.template_label.config(text=f"模板文件: {self.template_path}")

    def validate_date_entry(self, event) -> None:
        """驗證手動輸入的日期"""
        date_str = self.date_entry.get().strip()
        try:
            datetime.strptime(date_str, "%Y-%m-%d")
            self.selected_date = date_str
        except ValueError:
            messagebox.showwarning("警告", "日期格式錯誤，請使用 YYYY-MM-DD 格式，例如 2025-03-24。")
            self.date_entry.delete(0, tk.END)
            self.date_entry.insert(0, self.selected_date)

    def open_calendar(self) -> None:
        """打開日曆選擇器"""
        top = tk.Toplevel(self.root)
        top.title("選擇日期")
        cal = Calendar(top, selectmode="day", year=2025, month=3, day=24, date_pattern="yyyy-mm-dd")
        cal.pack(pady=10)
        def grab_date():
            self.selected_date = cal.get_date()
            self.date_entry.delete(0, tk.END)
            self.date_entry.insert(0, self.selected_date)
            top.destroy()
        tk.Button(top, text="確定", command=grab_date, bg="#4CAF50", fg="white", font=("微軟正黑體", 10)).pack(pady=5)

    def start_conversion(self) -> None:
        """開始轉換流程"""
        if not self.file_path:
            messagebox.showerror("錯誤", "請先選擇一個Excel文件。")
            return
        if not self.template_path:
            messagebox.showerror("錯誤", "請先選擇一個模板文件。")
            return

        date_str = self.date_entry.get().strip()
        if not date_str:
            messagebox.showerror("錯誤", "請選擇或輸入日期。")
            return

        try:
            datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError:
            messagebox.showerror("錯誤", "日期格式錯誤，請使用 YYYY-MM-DD 格式，例如 2025-03-24。")
            return

        self.selected_date = date_str
        self.output_text.delete(1.0, tk.END)
        process_excel_file(self.file_path, self.template_path, date_str, self.output_text)
        messagebox.showinfo("完成", "Word文件生成完成！")

# ============================
# 主程式入口
# ----------------------------

def main() -> None:
    root = tk.Tk()
    App(root)
    root.mainloop()

if __name__ == "__main__":
    main()